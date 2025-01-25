import sys
from PyQt5.QtCore import Qt, QEvent, QBuffer
from PyQt5.QtGui import QFont, QIcon
from PyQt5.QtWidgets import (
    QApplication,
    QMainWindow,
    QTextEdit,
    QFileDialog,
    QMenu,
    QAction
)
from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup
import base64
import io
from PIL import Image
import tempfile

class NoteApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.init_ui()
        self.current_file = None  # To track the file being saved/loaded

    def init_ui(self):
        self.setWindowTitle("GlassifyNotes")
        self.setWindowIcon(QIcon("note_icon.ico"))
        self.setGeometry(100, 100, 800, 500)
        self.setWindowFlags(Qt.FramelessWindowHint | Qt.WindowStaysOnTopHint)
        self.setAttribute(Qt.WA_TranslucentBackground)

        self.text_edit = QTextEdit(self)
        self.text_edit.setGeometry(0, 0, 800, 500)
        self.text_edit.setFont(QFont("Segoe UI", 14))
        self.text_edit.setStyleSheet(
            "background-color: white; border-radius: 10px; padding: 10px;"
        )
        self.text_edit.installEventFilter(self)

        self.text_edit.setContextMenuPolicy(Qt.CustomContextMenu)
        self.text_edit.customContextMenuRequested.connect(self.show_context_menu)

        self.text_edit.setFocus()
        self.text_edit.setAcceptRichText(True)
        self.text_edit.setPlainText("")
        self.text_edit.keyPressEvent = self.custom_key_press

        self.set_transparent()

    def eventFilter(self, source, event):
        if source == self.text_edit:
            if event.type() == QEvent.Enter:
                self.set_active()
            elif event.type() == QEvent.Leave:
                self.set_transparent()
        return super().eventFilter(source, event)

    def set_active(self):
        self.setWindowOpacity(0.75)
        self.text_edit.setStyleSheet(
            "background-color: white; border-radius: 10px; padding: 10px;"
        )

    def set_transparent(self):
        self.setWindowOpacity(0.3)
        self.text_edit.setStyleSheet(
            "background-color: rgba(255, 255, 255, 0.5); border-radius: 10px; padding: 10px;"
        )

    def custom_key_press(self, event):
        if event.key() == Qt.Key_S and event.modifiers() == Qt.ControlModifier:
            self.save_note()
        elif (
            event.key() == Qt.Key_S
            and event.modifiers() == (Qt.ControlModifier | Qt.ShiftModifier)
        ):
            self.save_as_note()
        elif event.key() == Qt.Key_Q and event.modifiers() == Qt.ControlModifier:
            self.save_and_exit()
        elif event.key() == Qt.Key_V and event.modifiers() == Qt.ControlModifier:
            self.paste_content()
        else:
            QTextEdit.keyPressEvent(self.text_edit, event)

    def save_note(self):
        """Save the note if it has already been saved, otherwise prompt for Save As."""
        if self.current_file:
            self.save_to_file(self.current_file)
        else:
            self.save_as_note()

    def save_as_note(self):
        options = QFileDialog.Options()
        file_name, _ = QFileDialog.getSaveFileName(
            self,
            "Save Note As",
            "",
            "Word Document (*.docx);;HTML Files (*.html);;All Files (*)",
            options=options,
        )
        if file_name:
            self.current_file = file_name
            self.save_to_file(file_name)

    def save_and_exit(self):
        if not self.current_file:
            self.save_as_note()
        else:
            self.save_to_file(self.current_file)
        QApplication.quit()

    def save_to_file(self, file_name):
        if file_name.endswith(".docx"):
            self.save_as_docx(file_name)
        else:
            with open(file_name, "w", encoding="utf-8") as file:
                file.write(self.text_edit.toHtml())

    def save_as_docx(self, file_name):
        """Save the note as a .docx file with images using tempfile."""
        
        document = Document()

        # Parse the HTML content
        html_content = self.text_edit.toHtml()
        soup = BeautifulSoup(html_content, "html.parser")

        for element in soup.body.descendants:
            if element.name == "p":
                document.add_paragraph(element.get_text())
            elif element.name == "img":
                # Extract base64 image data
                img_data = element["src"].split(",")[1]
                image_bytes = base64.b64decode(img_data)

                # Convert bytes to an image
                image_stream = io.BytesIO(image_bytes)
                image = Image.open(image_stream)

                # Use tempfile to create a temporary file in memory
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_image:
                    image.save(temp_image, format="PNG")
                    temp_image_path = temp_image.name

                # Add the image to the document
                document.add_picture(temp_image_path, width=Inches(4))  # Adjust size as needed

        # Save the document
        document.save(file_name)

    def show_context_menu(self, position):
        menu = QMenu()
        paste_action = QAction("Paste", self)
        paste_action.triggered.connect(self.paste_content)
        menu.addAction(paste_action)
        menu.exec_(self.text_edit.mapToGlobal(position))

    def paste_content(self):
        clipboard = QApplication.clipboard()
        mime_data = clipboard.mimeData()

        if mime_data.hasImage():
            image = clipboard.image()
            cursor = self.text_edit.textCursor()
            cursor.insertHtml(
                f'<img src="data:image/png;base64,{self.image_to_base64(image)}">'
            )
        elif mime_data.hasText():
            self.text_edit.paste()

    def image_to_base64(self, image):
        buffer = QBuffer()
        buffer.open(QBuffer.WriteOnly)
        image.save(buffer, "PNG")
        byte_array = buffer.data()
        return byte_array.toBase64().data().decode()

    def mousePressEvent(self, event):
        if event.button() == Qt.LeftButton:
            self.drag_pos = event.globalPos() - self.frameGeometry().topLeft()
            event.accept()

    def mouseMoveEvent(self, event):
        if event.buttons() == Qt.LeftButton:
            self.move(event.globalPos() - self.drag_pos)
            event.accept()

if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    note_app = NoteApp()
    note_app.show()
    sys.exit(app.exec_())
