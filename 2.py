import sys
import base64
import io
import tempfile
from PyQt5.QtCore import Qt, QEvent, QBuffer
from PyQt5.QtGui import QFont
from PyQt5.QtWidgets import (
    QApplication,
    QMainWindow,
    QTextEdit,
    QFileDialog,
    QMenu,
    QAction,
    QMessageBox
)
from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup
from PIL import Image


# -------------------- Model --------------------
class NoteModel:
    def save_to_file(self, content, file_name):
        """Handle file saving logic"""
        try:
            if file_name.endswith(".docx"):
                self._save_as_docx(content, file_name)
            else:
                with open(file_name, "w", encoding="utf-8") as file:
                    file.write(content)
            return True
        except Exception as e:
            raise Exception(f"Save failed: {str(e)}")

    def _save_as_docx(self, html_content, file_name):
        """Handle DOCX conversion logic"""
        try:
            document = Document()
            soup = BeautifulSoup(html_content, "html.parser")

            for element in soup.body.descendants:
                if element.name == "p":
                    document.add_paragraph(element.get_text())
                elif element.name == "img":
                    self._add_image_to_doc(document, element)

            document.save(file_name)
        except Exception as e:
            raise Exception(f"DOCX conversion failed: {str(e)}")

    def _add_image_to_doc(self, document, element):
        """Handle image processing for DOCX"""
        try:
            img_data = element["src"].split(",")[1]
            image_bytes = base64.b64decode(img_data)
            image_stream = io.BytesIO(image_bytes)
            image = Image.open(image_stream)

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_image:
                image.save(temp_image, format="PNG")
                document.add_picture(temp_image.name, width=Inches(4))
        except Exception as e:
            raise Exception(f"Image processing failed: {str(e)}")


# -------------------- View --------------------
class NoteView(QMainWindow):
    def __init__(self, controller):
        super().__init__()
        self.controller = controller
        self.init_ui()
        self.current_file = None

    def init_ui(self):
        self.setWindowTitle("Smart Note")
        self.setGeometry(100, 100, 800, 500)
        self.setWindowFlags(Qt.FramelessWindowHint | Qt.WindowStaysOnTopHint)
        self.setAttribute(Qt.WA_TranslucentBackground)

        self.text_edit = QTextEdit(self)
        self.text_edit.setGeometry(0, 0, 800, 500)
        self.text_edit.setFont(QFont("Segoe print", 14))
        self.text_edit.setStyleSheet(
            "background-color: white; border-radius: 10px; padding: 10px;"
        )
        self.text_edit.installEventFilter(self)

        self.text_edit.setContextMenuPolicy(Qt.CustomContextMenu)
        self.text_edit.customContextMenuRequested.connect(self.show_context_menu)

        self.text_edit.setFocus()
        self.text_edit.setAcceptRichText(True)
        self.set_transparent()

    # ... (Keep previous UI methods unchanged, add error display methods)

    def show_error(self, message):
        QMessageBox.critical(self, "Error", message)

    def show_context_menu(self, position):
        menu = QMenu()
        
        undo_action = QAction("Undo", self)
        undo_action.triggered.connect(self.text_edit.undo)
        undo_action.setEnabled(self.text_edit.document().isUndoAvailable())
        
        redo_action = QAction("Redo", self)
        redo_action.triggered.connect(self.text_edit.redo)
        redo_action.setEnabled(self.text_edit.document().isRedoAvailable())
        
        paste_action = QAction("Paste", self)
        paste_action.triggered.connect(self.controller.paste_content)

        menu.addAction(undo_action)
        menu.addAction(redo_action)
        menu.addAction(paste_action)
        menu.exec_(self.text_edit.mapToGlobal(position))


# -------------------- Controller --------------------
class NoteController:
    def __init__(self):
        self.model = NoteModel()
        self.view = NoteView(self)
        self.connect_signals()

    def connect_signals(self):
        self.view.text_edit.textChanged.connect(self.enable_undo_redo)

    def enable_undo_redo(self):
        """Update undo/redo availability"""
        # Can be extended to update UI indicators
        pass

    def custom_key_press(self, event):
        if event.key() == Qt.Key_S and event.modifiers() == Qt.ControlModifier:
            self.save_note()
        elif event.key() == Qt.Key_S and event.modifiers() == (Qt.ControlModifier | Qt.ShiftModifier):
            self.save_as_note()
        elif event.key() == Qt.Key_Q and event.modifiers() == Qt.ControlModifier:
            self.save_and_exit()
        elif event.key() == Qt.Key_Z and event.modifiers() == Qt.ControlModifier:
            self.view.text_edit.undo()
        elif event.key() == Qt.Key_Y and event.modifiers() == Qt.ControlModifier:
            self.view.text_edit.redo()
        else:
            QTextEdit.keyPressEvent(self.view.text_edit, event)

    def save_note(self):
        if self.view.current_file:
            try:
                success = self.model.save_to_file(
                    self.view.text_edit.toHtml(),
                    self.view.current_file
                )
                if not success:
                    self.view.show_error("Save operation failed")
            except Exception as e:
                self.view.show_error(str(e))
        else:
            self.save_as_note()

    def save_as_note(self):
        options = QFileDialog.Options()
        file_name, _ = QFileDialog.getSaveFileName(
            self.view,
            "Save Note As",
            "",
            "Word Document (*.docx);;HTML Files (*.html);;All Files (*)",
            options=options
        )
        if file_name:
            try:
                success = self.model.save_to_file(
                    self.view.text_edit.toHtml(),
                    file_name
                )
                if success:
                    self.view.current_file = file_name
                else:
                    self.view.show_error("Save operation failed")
            except Exception as e:
                self.view.show_error(str(e))

    # ... (Keep other controller methods with error handling)


if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    controller = NoteController()
    controller.view.show()
    sys.exit(app.exec_())